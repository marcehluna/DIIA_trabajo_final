"""Genera el documento Word con el resumen de la implementación (entrega / evaluador)."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING


ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "docs" / "Resumen-implementacion-Asistente-Regatas.docx"


def _set_body_style(document: Document) -> None:
    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    p_fmt = style.paragraph_format
    p_fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p_fmt.space_after = Pt(6)


def _h(doc: Document, text: str, level: int) -> None:
    doc.add_heading(text, level=level)


def _p(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def _bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def _tech_table(doc: Document) -> None:
    rows = [
        (
            "Python 3",
            "Lenguaje de implementación del asistente, orquestación del pipeline y configuración vía entorno.",
        ),
        (
            "Gradio",
            "Framework de interfaz web (formularios, desplegables, salida de texto) para probar el flujo sin desarrollar front a medida. Pensado para demos locales y publicación en Hugging Face Spaces.",
        ),
        (
            "pypdf",
            "Lectura y extracción de texto desde los PDF del corpus (Call Book y Case Book) para su partición en fragmentos (chunks).",
        ),
        (
            "OpenAI SDK (paquete openai)",
            "Cliente HTTP hacia API compatible con el estándar de OpenAI: chat completions (`/v1/chat/completions`) y, si aplica, embeddings en la ruta que exponga el proveedor.",
        ),
        (
            "Ollama (u otro host compatible)",
            "Servidor local típico que expone un endpoint tipo OpenAI; permite ejecutar modelos (p. ej. Llama) sin depender de un servicio comercial, manteniendo la misma abstracción de código `HTTPChatClient`.",
        ),
        (
            "RAG (Retrieval-Augmented Generation)",
            "Patrón: antes de generar la respuesta, se recupera contexto del corpus; el modelo recibe el relato del usuario y fragmentos relevantes, reduciendo alucinaciones y anclando el informe a las reglas oficiales.",
        ),
        (
            "Recuperación léxica (BM25-like / solapamiento)",
            "Implementación por defecto que no requiere embeddings: rankea fragmentos según términos, ligera y adecuada a entornos sin GPU o sin API de embeddings.",
        ),
        (
            "Recuperación semántica (opcional)",
            "Vectores y similitud coseno, vía API HTTP o modelo local (sentence-transformers) según `REGATAS_EMBEDDING_BACKEND` y configuración asociada.",
        ),
        (
            "NumPy",
            "Operaciones numéricas usadas en la pila de embeddings/retrieval cuando corresponde.",
        ),
        (
            "Hugging Face Spaces (opcional)",
            "Plataforma de alojamiento descrita en el README: por defecto puede usarse un backend `stub` sin LLM remoto, o conectar API y secretos para un modelo real en la nube.",
        ),
        (
            "Variables de entorno / `.env`",
            "Desacoplan ruta al corpus, modo LLM, URLs, claves, idioma y estrategia de prompt, permitiendo el mismo código en distintos despliegues (local, Space, contenedor).",
        ),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Tecnología o concepto"
    hdr[1].text = "Descripción breve"
    for t, d in rows:
        row = table.add_row().cells
        row[0].text = t
        row[1].text = d


def _comp_table(doc: Document) -> None:
    rows = [
        (
            "app.py",
            "Interfaz Gradio: carga (o reutiliza) el pipeline, campos de relato de quien protesta y, opcionalmente, del protestado, elección de idioma del system prompt, estrategia de prompt, modelo (cuando aplica) y presentación del informe generado.",
        ),
        (
            "regatas_assistant/config.py (Settings)",
            "Carga de configuración centralizada: rutas y nombres del corpus, tamaño y solapamiento de chunks, `top_k` de recuperación, backend y modelo de LLM, backend de embeddings, idioma y estrategia de prompt. Uso: definir variables de entorno o confiar en valores por defecto; instanciar con `Settings.from_env()`.",
        ),
        (
            "regatas_assistant/pipeline.py (ProtestPipeline)",
            "Orquesta el análisis: compone la consulta RAG, recupera fragmentos, formatea el contexto, arregla el system/user prompt y delega la generación al `LLMClient`. Uso: `ProtestPipeline.from_env()` o `from_settings` tras construir `Settings` en tests.",
        ),
        (
            "regatas_assistant/ingestion.py",
            "Ingesta PDFs en `TextChunk` (texto particionado) y formatea los fragmentos seleccionados para inyectarlos en el prompt. No expone lógica de reglas: solo estructuración y lectura de documentos.",
        ),
        (
            "regatas_assistant/rag/ (retriever.py y backends)",
            "`build_retriever` y `CorpusRetriever` encapsulan el modo léxico, HTTP o local. Sirven para sustituir la estrategia de búsqueda sin cambiar el resto de la app.",
        ),
        (
            "regatas_assistant/prompts.py",
            "Plantillas del system prompt y del usuario, normalización de idioma (es/en) y estrategia (p. ej. CoT y variantes) para alinear el tono y estructura del informe con el tramo académico (PLN).",
        ),
        (
            "regatas_assistant/llm/ (base, HTTPChatClient, stub)",
            "Interfaz común `LLMClient`. `HTTPChatClient` habla con la API compatible OpenAI. `StubLLMClient` devuelve salida fija: útil para demos o Spaces sin clave, para validar el flujo de UI y RAG sin generación real.",
        ),
        (
            "regatas_assistant/ollama_models.py",
            "Consulta a la API de Ollama (`/api/tags`) para listar modelos instalados cuando se usa Ollama como base URL, poblando el selector en la interfaz.",
        ),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Componente (archivo o paquete)"
    hdr[1].text = "Función y uso"
    for c, f in rows:
        row = table.add_row().cells
        row[0].text = c
        row[1].text = f


def build() -> Path:
    document = Document()
    _set_body_style(document)
    # Portada
    t = document.add_paragraph()
    t.add_run("Resumen de implementación\n").bold = True
    t.add_run("Asistente de protestas — Team Racing (PoC)\n").bold = True
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for _ in range(2):
        document.add_paragraph()

    _h(document, "1. Resumen del trabajo", 1)
    _p(
        document,
        "Se implementó un asistente de apoyo a la redacción de análisis de "
        "protestas en regatas de team racing, orientado a docentes o evaluadores. "
        "El sistema aplica técnicas de Procesamiento del Lenguaje Natural en un "
        "flujo RAG: a partir de relatos en lenguaje natural, recupera fragmentos de "
        "un corpus normativo (PDF) y pide a un modelo de lenguaje que elabore un informe "
        "estructurado, con instrucciones definidas en prompts configurables. La "
        "arquitectura desacopla la interfaz, la configuración, la ingesta, la "
        "recuperación, los prompts y el motor de lenguaje, de modo que el prototipo se "
        "puede ejecutar de forma local (Ollama) o desplegarse con variantes mínimas de configuración."
    )

    _h(document, "2. Tecnologías y conceptos", 1)
    _tech_table(document)

    _h(document, "3. Componentes arquitecturales", 1)
    _p(
        document,
        "A continuación se resumen los bloques de software implementados, su finalidad y cómo encajan en el uso del sistema (usuario final o desarrollador)."
    )
    _comp_table(document)

    _h(document, "4. Flujo de datos (visión de proceso)", 1)
    for line in [
        "Arranque: se lee la configuración, se indexan o cargan en memoria los fragmentos del corpus PDF, se instancia el retriever y el cliente LLM según el backend elegido.",
        "Consulta: el usuario describe la protesta (y, opcionalmente, la versión del barco protestado) en la interfaz.",
        "RAG: el sistema construye una consulta de búsqueda, recupera los k fragmentos más relevantes (léxico o semántico, según configuración) y los formatea como contexto legible en el prompt.",
        "Generación: se arman mensajes con system prompt (idioma/estrategia) y user prompt con el contexto recuperado y los relatos; el LLM produce el informe, que se muestra en la UI.",
    ]:
        _bullet(document, line)

    _h(document, "5. Cómo utilizarlo (indicaciones prácticas)", 1)
    for line in [
        "Requisito de activos: en la raíz del proyecto o en la ruta configurada deben estar los PDF del Call Book y Case Book indicados en la documentación o en `REGATAS_CORPUS_FILES`.",
        "Entorno local recomendado: `pip install -r requirements.txt`, Ollama en marcha (p. ej. en el puerto 11434) y modelo descargado (`ollama pull …`), luego `python app.py` abre la interfaz web.",
        "Parametrización: variables como `REGATAS_LLM_BACKEND` (http o stub), `REGATAS_LLM_BASE_URL`, `REGATAS_LLM_MODEL`, `REGATAS_EMBEDDING_BACKEND` (lexical, http, local) controlan dónde corre el modelo y cómo se hace la búsqueda, sin tocar el código de negocio.",
    ]:
        _bullet(document, line)

    _h(document, "6. Detalles relevantes para el evaluador", 1)
    for line in [
        "El alcance es una prueba de concepto (PoC): demuestra integración ingesta RAG, prompts y despliegue, no reemplaza el criterio arbitral ni la validez jurídica de una protesta en la vida real.",
        "La trazabilidad hacia las fuentes mejora si se explicitan en el diseño de prompts citas o referencia a fragmentos recuperados; el corpus oficial es el ancla del conocimiento.",
        "Modularidad: añadir un proveedor de LLM compatible OpenAI, cambiar el retriever a semántico local, o ajustar el tono del system prompt son tareas acotadas gracias a `Settings` y a las abstracciones de `llm` y `rag`.",
    ]:
        _bullet(document, line)

    document.add_paragraph()
    _p(
        document,
        "Documento generado de forma asistida a partir de la estructura del repositorio y de la arquitectura documentada, para acompañar la exposición académica del trabajo de PLN."
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(OUT))
    return OUT


if __name__ == "__main__":
    p = build()
    print(f"Generado: {p}")
